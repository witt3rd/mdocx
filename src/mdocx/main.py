import argparse
import json

from docx import Document
from mistletoe import Document as MdDocument
from mistletoe.ast_renderer import ASTRenderer

#


heading_styles = [
    "H1 - Chapter",
    "H1 - Section",
    "H2 - Heading",
    "H3 - Subheading",
    "H4 - Subheading",
    "H5 - Subheading",
    "H6 - Subheading",
]

style_types = {
    "H1 - Chapter": "PARAGRAPH",
    "Default Paragraph Font": "CHARACTER",
    "P - Regular": "PARAGRAPH",
    "L - Bullets": "PARAGRAPH",
    "H1 - Section": "PARAGRAPH",
    "H2 - Heading": "PARAGRAPH",
    "P - Code": "CHARACTER",
    "SC - Source": "PARAGRAPH",
    "SC - Highlight": "CHARACTER",
    "P - Callout Heading": "PARAGRAPH",
    "P - Callout": "PARAGRAPH",
    "P - Source": "PARAGRAPH",
    "L - Numbers": "PARAGRAPH",
    "L - Regular": "PARAGRAPH",
    "L - Source": "PARAGRAPH",
    "P - Bold": "CHARACTER",
    "P - Italics": "CHARACTER",
    "P - Keyword": "CHARACTER",
    "P - URL": "CHARACTER",
    "Normal": "PARAGRAPH",
    "SC - Heading": "PARAGRAPH",
    "SC - Link": "PARAGRAPH",
    "IMG - Caption": "PARAGRAPH",
    "H3 - Subheading": "PARAGRAPH",
    "H4 - Subheading": "PARAGRAPH",
    "P - Quote": "PARAGRAPH",
    "SP - Editorial": "PARAGRAPH",
}


#


def get_style_type(style_name):
    return style_types.get(style_name, "PARAGRAPH")


def get_code_content(node):
    code_content = node.get("content", "")
    if not code_content and "children" in node:
        code_content = "".join(
            child["content"] for child in node["children"] if "content" in child
        )

    return code_content


def process_list_item(node, docx, list_leader, level=0):
    list_style = "L - Bullets" if list_leader == "-" else "L - Numbers"
    current_paragraph = docx.add_paragraph(style=list_style)
    current_paragraph.clear()

    for child in node["children"]:
        if child["type"] == "Paragraph":
            process_list_item_content(child, current_paragraph)
        elif child["type"] == "List":
            for nested_item in child["children"]:
                process_list_item(nested_item, docx, child.get("leader"), level + 1)
        else:
            process_node(child, docx, current_paragraph)


def process_list_item_content(node, paragraph):
    for child in node["children"]:
        if child["type"] == "InlineCode":
            code_content = get_code_content(child)
            run = paragraph.add_run(code_content)
            run.style = "P - Code"
        elif child["type"] == "Strong":
            process_strong(child, paragraph)
        elif child["type"] == "Emphasis":
            process_emphasis(child, paragraph)
        elif child["type"] == "RawText":
            paragraph.add_run(child["content"])
        else:
            process_node(child, None, paragraph)


def process_inline_style(node, paragraph):
    if node["type"] == "Strong":
        run = paragraph.add_run()
        process_strong(node, paragraph)
    elif node["type"] == "Emphasis":
        run = paragraph.add_run()
        process_emphasis(node, paragraph)
    elif node["type"] == "InlineCode":
        run = paragraph.add_run(node.get("content", ""))
        run.style = "P - Code"
    else:
        process_node(node, None, paragraph)


def process_strong(node, paragraph):
    for child in node["children"]:
        if child["type"] == "RawText":
            run = paragraph.add_run(child["content"])
            run.bold = True
        elif child["type"] == "Emphasis":
            process_emphasis(child, paragraph, bold=True)
        else:
            process_inline_style(child, paragraph)


def process_emphasis(node, paragraph, bold=False):
    for child in node["children"]:
        if child["type"] == "RawText":
            run = paragraph.add_run(child["content"])
            run.italic = True
            if bold:
                run.bold = True
        elif child["type"] == "Strong":
            process_strong(child, paragraph)
            paragraph.runs[-1].italic = True
        else:
            process_inline_style(child, paragraph)


def process_node(node, docx, current_paragraph=None):
    node_type = node["type"]

    if node_type == "Document":
        for child in node["children"]:
            process_node(child, docx)

    elif node_type == "Heading":
        heading_style = heading_styles[node["level"] - 1]
        current_paragraph = docx.add_paragraph(style=heading_style)
        for child in node["children"]:
            process_node(child, docx, current_paragraph)

    elif node_type == "Paragraph":
        current_paragraph = docx.add_paragraph(style="P - Regular")
        for child in node["children"]:
            process_node(child, docx, current_paragraph)

    elif node_type == "RawText":
        if current_paragraph:
            if current_paragraph.style.name in ["L - Bullets", "L - Numbers"]:
                # For list items, ensure we're not adding extra content at the start
                if not current_paragraph.runs:
                    current_paragraph.add_run(node["content"])
                else:
                    current_paragraph.runs[-1].add_text(node["content"])
            else:
                current_paragraph.add_run(node["content"])
        else:
            docx.add_paragraph(node["content"], style="P - Regular")

    elif node_type == "List":
        for child in node["children"]:
            process_list_item(child, docx, child.get("leader"))

    elif node_type == "ListItem":
        list_style = "L - Bullets" if node.get("leader") == "-" else "L - Numbers"
        current_paragraph = docx.add_paragraph(style=list_style)
        for child in node["children"]:
            if child["type"] == "Paragraph":
                for subchild in child["children"]:
                    process_node(subchild, docx, current_paragraph)
            else:
                process_node(child, docx, current_paragraph)

    elif node_type == "InlineCode":
        code_content = get_code_content(node)
        if current_paragraph:
            current_paragraph.add_run(code_content).style = "P - Code"
        else:
            p = docx.add_paragraph(style="P - Regular")
            p.add_run(code_content).style = "P - Code"

    elif node_type == "CodeFence":
        p = docx.add_paragraph(style="P - Source")
        code_content = node.get("content", "")
        if not code_content and "children" in node:
            code_content = "".join(
                child.get("content", "") for child in node["children"]
            ).strip()
        p.add_run(code_content)

    elif node_type == "Strong":
        if current_paragraph:
            process_strong(node, current_paragraph)
        else:
            p = docx.add_paragraph(style="P - Regular")
            process_strong(node, p)

    elif node_type == "Emphasis":
        if current_paragraph:
            process_emphasis(node, current_paragraph)
        else:
            p = docx.add_paragraph(style="P - Regular")
            process_emphasis(node, p)

    elif node_type == "LineBreak":
        if current_paragraph:
            current_paragraph.add_run("\n")
        else:
            docx.add_paragraph()

    elif node_type == "Link":
        link_text = "".join(
            child["content"] for child in node["children"] if "content" in child
        )
        link_url = node.get("target", "")
        if current_paragraph:
            current_paragraph.add_run(link_text)
            current_paragraph.add_run(" (")
            url_run = current_paragraph.add_run(link_url)
            url_run.style = "P - URL"
            current_paragraph.add_run(")")
        else:
            p = docx.add_paragraph(style="P - Regular")
            p.add_run(link_text)
            p.add_run(" (")
            url_run = p.add_run(link_url)
            url_run.style = "P - URL"
            p.add_run(")")
    else:
        print(f"Unhandled node type: {node_type}")


def markdown_to_docx(md_file, docx_template, output_file):
    with open(md_file, "r", encoding="utf-8") as fin:
        with ASTRenderer() as renderer:
            docx = Document(docx_template)
            mdoc = MdDocument(fin)
            ast = json.loads(renderer.render(mdoc))
            process_node(ast, docx)
            docx.save(output_file)


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX with custom styling."
    )
    parser.add_argument("input_markdown", help="Path to the input Markdown file")
    parser.add_argument("template_docx", help="Path to the template DOCX file")
    parser.add_argument("output_docx", help="Path for the output DOCX file")

    args = parser.parse_args()

    try:
        markdown_to_docx(args.input_markdown, args.template_docx, args.output_docx)
        print(f"Successfully converted {args.input_markdown} to {args.output_docx}")
    except Exception as e:
        print(f"An error occurred: {str(e)}")


if __name__ == "__main__":
    main()
